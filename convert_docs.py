import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def make_proje_raporu_docx():
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Yapay Zeka ve Veri Bilimi Proje Raporu\nAkıllı Şehir İçi Trafik ve Toplu Taşıma Optimize Sistemi (AI-Trafik)")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(0, 102, 153)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Social Office — Yapay Zeka Mesleki Gelişim Programı (Ödev 4)")
    run_sub.font.size = Pt(11)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    doc.add_heading("1. Giriş ve Problem Tanımı", level=1)
    doc.add_paragraph("Günümüz büyükşehirlerinde yaşanan en büyük günlük sorunlardan biri şehir içi trafik sıkışıklığıdır. Mevcut sinyalizasyon sistemleri sabit zaman aralıklarına göre çalıştığı için bir şeritte araç yokken diğer şeritte kilometrelerce kuyruk oluşabilmektedir.\n\n• Etkilenen Kitle: Sürücüler, toplu taşıma yolcuları, lojistik ve acil durum araçları.\n• Temel Nedenler: Sabit zamanlı ışıklar, anlık kaza/yol çalışması tepkisizliği.\n• Çözümün Önemi: Zaman kaybını engellemek, rölanti yakıt tüketimini ve karbon salınımını %10-25 azaltmak.")

    doc.add_heading("2. Veri ve Analiz", level=1)
    doc.add_paragraph("• İhtiyaç Duyulan Veriler: Trafik kameraları canlı görüntüleri, anlık bekleme süreleri ve geçmiş yoğunluk verileri.\n• Veri Kaynakları: Belediye Ulaşım Kontrol Merkezleri (UKM) kameraları ve anonim GPS verileri.\n• Gizlilik ve Etik: Araç plakaları ve sürücü yüzleri işlendiği anda otomatik olarak bulanıklaştırılacaktır (KVKK uyumlu).")

    base_dir = os.path.dirname(os.path.abspath(__file__))
    g1 = os.path.join(base_dir, "gorseller", "gorsel_1_trafik_analiz_paneli.png")
    g2 = os.path.join(base_dir, "gorseller", "gorsel_2_ai_sinyalizasyon_akisi.png")
    g3 = os.path.join(base_dir, "gorseller", "gorsel_3_toplu_tasima_otomasyonu.png")

    if os.path.exists(g1):
        doc.add_paragraph("Şekil 1: Akıllı Şehir Trafik Kontrol Paneli Görseli").runs[0].font.italic = True
        doc.add_picture(g1, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3. Çözüm Önerileri Geliştirme", level=1)
    doc.add_paragraph("Geliştirilecek AI-Trafik sistemi, kameralardan alınan canlı görüntüleri işleyerek şeritlerdeki araç sayısını anlık tespit eder. Yoğun şeride yeşil ışık süresini uzatır, boş şeridin kırmızı ışığını kısaltır.\n\n• Başarı Metrikleri: Kavşak bekleme süresinde %25 azalma, ortalama seyir hızında %20 artış.\n• Test Yöntemi: 2 pilot kavşakta 30 günlük simülasyon ve saha testi.")

    if os.path.exists(g2):
        doc.add_paragraph("Şekil 2: AI Trafik ve Dinamik Sinyalizasyon Akış Diyagramı").runs[0].font.italic = True
        doc.add_picture(g2, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4. Yapay Zeka Entegrasyonu", level=1)
    doc.add_paragraph("• YZ Yöntemleri: Bilgisayarlı Görü (YOLO / Araç Tespiti) ve Pekiştirmeli Öğrenme (Işık Süresi Optimizasyonu).\n• Tepki Süresi: 50 milisaniye (Tam gerçek zamanlı çalışma).\n• Manuel Yöntemlere Üstünlüğü: 7/24 kesintisiz ve insan hatasından uzak otomatik karar mekanizması.")

    if os.path.exists(g3):
        doc.add_paragraph("Şekil 3: Toplu Taşıma ve Şehir İçi Trafik Optimizasyonu").runs[0].font.italic = True
        doc.add_picture(g3, width=Inches(5.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("5. Sonuç ve Öneriler", level=1)
    doc.add_paragraph("Önerilen AI-Trafik sistemi bireysel düzeyde günlük 20-30 dakika zaman tasarrufu sağlarken, toplumsal düzeyde şehir içi hava kirliliğini düşürecek ve acil araçlara hızlı geçiş üstünlüğü tanıyacaktır.")

    out_path = os.path.join(base_dir, "proje_raporu.docx")
    doc.save(out_path)
    print(f"[OK] proje_raporu.docx oluşturuldu: {out_path}")

def make_sunum_docx():
    doc = docx.Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PROJE SUNUMU (5 SLAYT)\nAkıllı Şehir İçi Trafik ve Toplu Taşıma Optimize Sistemi")
    run_title.bold = True
    run_title.font.size = Pt(18)
    run_title.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_heading("SLAYT 1: Giriş ve Problem Tanımı", level=1)
    doc.add_paragraph("• Problem: Şehir içi trafik sıkışıklığı ve sabit zamanlı trafik lambaları.\n• Amaç: Yapay zeka ile bekleme sürelerini %25 azaltmak.")

    doc.add_heading("SLAYT 2: Veri ve Analiz", level=1)
    doc.add_paragraph("• Kaynaklar: Trafik kameraları, anlık araç sayıları ve GPS verileri.\n• Gizlilik: KVKK uyumlu anında araç plakası ve yüz anonimleştirme.")

    doc.add_heading("SLAYT 3: Çözüm Önerisi (AI-Trafik)", level=1)
    doc.add_paragraph("• Mantık: Kameralar anlık araç sayar, yoğun şeride yeşil ışık süresi uzatılır.\n• Avantajlar: %25-30 bekleme süresi düşüşü, zaman ve yakıt tasarrufu.")

    doc.add_heading("SLAYT 4: Yapay Zeka Entegrasyonu", level=1)
    doc.add_paragraph("• Yöntemler: Bilgisayarlı Görü (YOLO) ve Pekiştirmeli Öğrenme.\n• Çalışma: 50 ms tepki süresiyle gerçek zamanlı entegrasyon.")

    doc.add_heading("SLAYT 5: Sonuç ve Toplumsal Fayda", level=1)
    doc.add_paragraph("• Bireysel Fayda: Günlük 20-30 dk zaman kazanımı.\n• Toplumsal Fayda: Temiz hava, düşük emisyon ve acil araç geçiş önceliği.")

    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sunum.docx")
    doc.save(out_path)
    print(f"[OK] sunum.docx oluşturuldu: {out_path}")

if __name__ == "__main__":
    make_proje_raporu_docx()
    make_sunum_docx()
